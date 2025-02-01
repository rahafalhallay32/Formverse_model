import json
import os
import openai
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from collections import Counter
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv  # For loading the API key from environment variables

# Load OpenAI API key from .env
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Function to load survey data from JSON
def load_survey_data(json_data):
    try:
        survey_data = json.loads(json_data)
        return survey_data
    except json.JSONDecodeError as e:
        print("Error loading JSON:", e)
        return None

# Function to dynamically identify closed-ended questions and analyze responses
def analyze_closed_end_questions(survey_data, doc):
    closed_end_questions = {}
    
    # Iterate over questions and their answers
    for question in survey_data['questions']:
        question_text = question['question']
        question_type = question.get('questionType', '').lower()  # Get question type (e.g., MCQ, True/False)
        answers = question.get('answers', [])  # Get answers or default to an empty list
        
        # Process answers for closed-ended questions
        for answer in answers:
            if isinstance(answer, list):  # Ensure answer is a list
                for a in answer:
                    a_lower = a.lower()
                    
                    # Handle Yes/No questions
                    if a_lower in ['yes', 'no']:
                        if question_text not in closed_end_questions:
                            closed_end_questions[question_text] = {'type': 'yes_no', 'Yes': 0, 'No': 0}
                        if a_lower == 'yes':
                            closed_end_questions[question_text]['Yes'] += 1
                        elif a_lower == 'no':
                            closed_end_questions[question_text]['No'] += 1
                    
                    # Handle True/False questions
                    elif a_lower in ['true', 'false']:
                        if question_text not in closed_end_questions:
                            closed_end_questions[question_text] = {'type': 'true_false', 'True': 0, 'False': 0}
                        if a_lower == 'true':
                            closed_end_questions[question_text]['True'] += 1
                        elif a_lower == 'false':
                            closed_end_questions[question_text]['False'] += 1
                    
                    # Handle MCQ questions
                    elif question_type == 'mcq':
                        if question_text not in closed_end_questions:
                            closed_end_questions[question_text] = {'type': 'mcq'}
                        if a_lower not in closed_end_questions[question_text]:
                            closed_end_questions[question_text][a_lower] = 0
                        closed_end_questions[question_text][a_lower] += 1

    # Generate charts for closed-ended questions
    for i, (question, data) in enumerate(closed_end_questions.items()):
        chart_filename = f"chart_{i}.png"
        
        if data['type'] == 'yes_no' or data['type'] == 'true_false':
            # Pie chart for Yes/No and True/False questions
            labels = list(data.keys())[1:]  # Exclude 'type' key
            sizes = list(data.values())[1:]  # Exclude 'type' value
            
            plt.figure(figsize=(6,6))
            plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors)
            plt.title(f"Response Distribution - {question}")
            plt.savefig(chart_filename)
            plt.close()
        
        elif data['type'] == 'mcq':
            # Bar chart for MCQ questions
            labels = list(data.keys())[1:]  # Exclude 'type' key
            counts = list(data.values())[1:]  # Exclude 'type' value
            
            plt.figure(figsize=(8,6))
            plt.bar(labels, counts, color=plt.cm.Paired.colors)
            plt.title(f"Response Distribution - {question}")
            plt.xlabel("Options")
            plt.ylabel("Count")
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(chart_filename)
            plt.close()

        # Add the chart to the Word document
        doc.add_paragraph(f"Question: {question}")
        doc.add_picture(chart_filename, width=Inches(4.0))
        doc.add_paragraph("\n")
        
        os.remove(chart_filename)

def generate_wordcloud_for_open_end(survey_data, doc):
    open_end_answers = []

    # Iterate over questions and their answers
    for question in survey_data['questions']:
        answers = question.get('answers', [])  # Get answers or default to an empty list
        
        # Process answers for open-ended questions
        for answer in answers:
            if isinstance(answer, list):  # Ensure answer is a list
                for a in answer:
                    if a.lower() not in ['yes', 'no']:  # Exclude closed-ended answers
                        open_end_answers.append(a)

    # Generate word cloud for open-ended answers
    if open_end_answers:
        word_counts = Counter(" ".join(open_end_answers).split())
        most_common_words = dict(word_counts.most_common(10))
        
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(most_common_words)
        wordcloud_filename = "wordcloud.png"
        wordcloud.to_file(wordcloud_filename)

        doc.add_paragraph("Top 10 Most Common Words in Open-Ended Questions")
        doc.add_picture(wordcloud_filename, width=Inches(4.0))
        doc.add_paragraph("\n")

        os.remove(wordcloud_filename)

def generate_analysis_with_gpt(survey_data):
    goal = survey_data['goal']
    hypothesis = survey_data['hypothesis']
    target_group = survey_data['targetGroup']
    time_taken = survey_data['timeTaken']
    
    formatted_responses = []
    
    # Iterate over questions and their answers
    for question in survey_data['questions']:
        question_text = question['question']
        answers = question.get('answers', [])  # Get answers or default to an empty list
        
        if answers:
            formatted_responses.append(f"Question: {question_text}")
            for answer in answers:
                formatted_responses.append(f"- Answer: {answer}")

    analyze_prompt = f"""
    You are a skilled data analyst tasked with analyzing survey responses containing both closed and open-ended questions. Your main goal is to deliver a structured and data-driven analysis with proper charts and thematic analysis.

    ### Research Details:
    - *Goal:* {goal}
    - *Hypothesis:* {hypothesis}
    - *Target Group:* {target_group}
    - *Time Taken (in minutes):* {time_taken}
    - *Survey Responses:* {', '.join(formatted_responses)}

    ### Expected Output:
    1. Provide a summary of the survey results.
    2. For closed-ended questions, analyze response distribution.
    3. Perform thematic analysis for open-ended questions.
    4. Evaluate the hypothesis based on the data determine if the hypothesis is: *Supported* or *Partially Supported* or *Not Supported*
    """
    
    response = openai.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": analyze_prompt}],
        max_tokens=700,
        temperature=0.5
    )
    
    return response.choices[0].message.content

def create_word_document(survey_data, gpt_analysis, output_path):
    # Ensure the output directory exists
    if not os.path.exists(output_path):
        os.makedirs(output_path)

    doc = Document()
    
    doc.add_heading('Survey Analysis Report', 0)
    doc.add_paragraph(f"Goal: {survey_data['goal']}")
    doc.add_paragraph(f"Hypothesis: {survey_data['hypothesis']}")
    doc.add_paragraph(f"Target Group: {survey_data['targetGroup']}")
    doc.add_paragraph(f"Time Taken (in minutes): {survey_data['timeTaken']}")
    
    doc.add_heading('Analysis:', level=1)
    doc.add_paragraph(gpt_analysis)
    
    analyze_closed_end_questions(survey_data, doc)
    generate_wordcloud_for_open_end(survey_data, doc)
    
    output_filename = os.path.join(output_path, "survey_analysis_report.docx")
    doc.save(output_filename)
    print(f"Analysis saved to {output_filename}")
    
    return output_filename

def run_analysis(json_data, output_path):
    survey_data = load_survey_data(json_data)
    if survey_data:
        gpt_analysis = generate_analysis_with_gpt(survey_data)
        return create_word_document(survey_data, gpt_analysis, output_path)